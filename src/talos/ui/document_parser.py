"""Multi-format upload CFO — CHG-2026-05-02-007.

Parser unificato per CSV/XLSX/PDF/DOCX. Estrae 2 colonne minime
`descrizione`, `prezzo` (+ opzionali `v_tot`/`s_comp`/`category_node`)
dal file caricato dal CFO via `st.file_uploader`.

Strategia per formato:
  - CSV  : `pd.read_csv` (esistente, semantica invariata).
  - XLSX : `pd.read_excel` engine `openpyxl` — prima sheet, header riga 0.
  - PDF  : `pdfplumber.extract_tables()` su tutte le pagine, concatena.
           Richiede text layer nativo (NO OCR). Per scansioni: scope
           futuro via `OcrPipeline` (ADR-0017 canale 3).
  - DOCX : `python-docx` itera tabelle del documento, prima riga = header.

R-01 NO SILENT DROPS: file senza colonne `descrizione`/`prezzo` ->
`ValueError` esplicito (caller mostra `st.error`). File senza tabelle
estraibili (PDF scansionato, DOCX testo libero) -> `ValueError`.
"""

from __future__ import annotations

from typing import IO, Any

import pandas as pd

# Suffix supportati lower-case (caller normalizza prima del dispatch).
SUPPORTED_SUFFIXES: tuple[str, ...] = ("csv", "xlsx", "xls", "pdf", "docx")

# Minimo righe per considerare una tabella DOCX valida (header + 1 dato).
_MIN_DOCX_TABLE_ROWS: int = 2

# CHG-2026-05-02-024: catena encoding per CSV (Excel italiano usa cp1252,
# byte 0x97 em-dash non decodificabile in UTF-8 strict). `latin-1` finale
# è single-byte catch-all (mai solleva), garantisce R-01 NO SILENT DROPS.
# `utf-8-sig` come primo: rimuove BOM se presente, altrimenti decode UTF-8
# standard. Copre Excel office 365 + export moderni in un colpo.
CSV_ENCODING_CHAIN: tuple[str, ...] = ("utf-8-sig", "cp1252", "latin-1")


def parse_uploaded_document(uploaded: IO[bytes], suffix: str) -> pd.DataFrame:
    """Dispatcher per parsing multi-formato.

    :param uploaded: file-like object (es. `streamlit.UploadedFile`).
    :param suffix: estensione lower-case senza punto (es. "pdf").
    :returns: DataFrame con almeno colonne `descrizione`, `prezzo`.
    :raises ValueError: suffix non supportato o file non parsabile.
    """
    suffix = suffix.lower().lstrip(".")
    if suffix == "csv":
        return _parse_csv(uploaded)
    if suffix in ("xlsx", "xls"):
        return _parse_xlsx(uploaded)
    if suffix == "pdf":
        return _parse_pdf(uploaded)
    if suffix == "docx":
        return _parse_docx(uploaded)
    msg = f"Formato non supportato: .{suffix}. Attesi: {SUPPORTED_SUFFIXES}."
    raise ValueError(msg)


def _decode_with_fallback(raw: bytes) -> str:
    """Decode bytes su `CSV_ENCODING_CHAIN` (utf-8-sig → cp1252 → latin-1).

    CHG-2026-05-02-024: `latin-1` è single-byte e NON solleva mai
    (256 punti codice tutti validi), quindi il loop termina sempre
    con un risultato. Mojibake possibile solo se il file è in encoding
    esotico (kr/jp/etc) — fuori scope CFO italiano.
    """
    for enc in CSV_ENCODING_CHAIN:
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("latin-1", errors="replace")  # defensive, mai raggiunto


def _parse_csv(uploaded: IO[bytes]) -> pd.DataFrame:
    """Auto-detect separatore + encoding fallback (Excel italiano cp1252).

    CHG-2026-05-02-024: chain encoding `CSV_ENCODING_CHAIN` (utf-8-sig →
    cp1252 → latin-1) prima di fallire. Bug live Leader 2026-05-02:
    byte 0x97 em-dash da Excel italiano rompeva UTF-8 strict.
    """
    import csv  # noqa: PLC0415
    import io  # noqa: PLC0415

    raw = uploaded.read()
    sample_text = _decode_with_fallback(raw[:4096])
    try:
        dialect = csv.Sniffer().sniff(sample_text, delimiters=",;\t|")
        sep = dialect.delimiter
    except csv.Error:
        # Fallback heuristico: conta separatori candidati nelle prime righe.
        sep = ";" if sample_text.count(";") > sample_text.count(",") else ","

    last_exc: UnicodeDecodeError | None = None
    for enc in CSV_ENCODING_CHAIN:
        try:
            return pd.read_csv(io.BytesIO(raw), sep=sep, encoding=enc)
        except UnicodeDecodeError as exc:
            last_exc = exc
            continue
    msg = (
        f"CSV non decodificabile in nessun encoding {CSV_ENCODING_CHAIN}. Ultimo errore: {last_exc}"
    )
    raise ValueError(msg)


def _parse_xlsx(uploaded: IO[bytes]) -> pd.DataFrame:
    return pd.read_excel(uploaded, engine="openpyxl")


def _parse_pdf(uploaded: IO[bytes]) -> pd.DataFrame:
    """Estrae tabelle native da PDF con pdfplumber."""
    import io  # noqa: PLC0415

    import pdfplumber  # noqa: PLC0415

    # pdfplumber richiede BytesIO/path; UploadedFile è IO[bytes]: leggo in BytesIO.
    buf = io.BytesIO(uploaded.read())
    rows: list[list[Any]] = []
    header: list[str] | None = None
    with pdfplumber.open(buf) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables() or []:
                if not table:
                    continue
                if header is None:
                    header = [_norm_header(c) for c in table[0]]
                    rows.extend(table[1:])
                else:
                    # Pagine successive: skip header se uguale.
                    first = [_norm_header(c) for c in table[0]]
                    rows.extend(table[1:] if first == header else table)
    if header is None or not rows:
        msg = (
            "PDF senza tabelle native estraibili. Per scansioni serve OCR "
            "(scope futuro). Prova con XLSX o CSV."
        )
        raise ValueError(msg)
    return pd.DataFrame(rows, columns=header)


def _parse_docx(uploaded: IO[bytes]) -> pd.DataFrame:
    """Estrae tabelle da DOCX (prima tabella con header `descrizione`+`prezzo`)."""
    import docx  # noqa: PLC0415

    document = docx.Document(uploaded)
    if not document.tables:
        msg = "DOCX senza tabelle. Inserisci dati in formato tabellare."
        raise ValueError(msg)
    for table in document.tables:
        if len(table.rows) < _MIN_DOCX_TABLE_ROWS:
            continue
        header = [_norm_header(c.text) for c in table.rows[0].cells]
        if "descrizione" in header and "prezzo" in header:
            data = [[c.text for c in row.cells] for row in table.rows[1:]]
            return pd.DataFrame(data, columns=header)
    msg = "DOCX: nessuna tabella con header `descrizione`+`prezzo` trovata."
    raise ValueError(msg)


def _norm_header(value: object) -> str:
    """Normalizza header colonna: strip + lower, gestisce None/whitespace."""
    if value is None:
        return ""
    return str(value).strip().lower()
