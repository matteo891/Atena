"""Unit test `talos.ui.document_parser` (CHG-2026-05-02-007)."""

from __future__ import annotations

import io

import pytest

from talos.ui.document_parser import SUPPORTED_SUFFIXES, parse_uploaded_document

pytestmark = pytest.mark.unit


def test_unsupported_suffix_raises() -> None:
    buf = io.BytesIO(b"x")
    with pytest.raises(ValueError, match="Formato non supportato"):
        parse_uploaded_document(buf, "txt")


def test_csv_dispatch() -> None:
    buf = io.BytesIO(b"descrizione,prezzo\nGalaxy S24,549\nGalaxy A54,329\n")
    df = parse_uploaded_document(buf, "csv")
    assert list(df.columns) == ["descrizione", "prezzo"]
    assert len(df) == 2


def test_csv_auto_detect_semicolon_excel_italiano() -> None:
    """CHG-016: CSV con `;` (Excel italiano) auto-detected."""
    buf = io.BytesIO(b"descrizione;prezzo\nGalaxy S24;549\nGalaxy A54;329\n")
    df = parse_uploaded_document(buf, "csv")
    assert list(df.columns) == ["descrizione", "prezzo"]
    assert len(df) == 2


def test_csv_auto_detect_tab_separator() -> None:
    """CHG-016: TSV-like (tab separator)."""
    buf = io.BytesIO(b"descrizione\tprezzo\nGalaxy S24\t549\nGalaxy A54\t329\n")
    df = parse_uploaded_document(buf, "csv")
    assert list(df.columns) == ["descrizione", "prezzo"]
    assert len(df) == 2


# ---------------------------------------------------------------------------
# CHG-2026-05-02-024: encoding fallback chain
# ---------------------------------------------------------------------------


def test_csv_utf8_with_italian_accents() -> None:
    """UTF-8 con accenti italiani (`è`/`à`/`ò`) → parsing OK."""
    raw = "descrizione,prezzo\nCaffè italiano,3.50\nGià pagato,5.00\n".encode()
    buf = io.BytesIO(raw)
    df = parse_uploaded_document(buf, "csv")
    assert list(df.columns) == ["descrizione", "prezzo"]
    assert df.iloc[0]["descrizione"] == "Caffè italiano"
    assert df.iloc[1]["descrizione"] == "Già pagato"


def test_csv_utf8_sig_bom() -> None:
    """UTF-8 con BOM (Excel office 365 export): `\\ufeff` prefix tollerato."""
    raw = "﻿descrizione,prezzo\nGalaxy S24,549\n".encode("utf-8-sig")
    buf = io.BytesIO(raw)
    df = parse_uploaded_document(buf, "csv")
    # Header pulito (senza BOM residuo grazie a utf-8-sig).
    assert "descrizione" in df.columns
    assert df.iloc[0]["descrizione"] == "Galaxy S24"


def test_csv_cp1252_em_dash_byte_0x97() -> None:
    """Bug live Leader 2026-05-02: byte 0x97 (em-dash cp1252) → fallback OK."""
    # 0x97 in cp1252 = "—" (em-dash). UTF-8 strict rifiuta.
    raw = (
        b"descrizione,prezzo\n"
        b"Samsung Galaxy S24 \x97 256GB Onyx,549\n"
        b"iPhone 15 Pro \x97 128GB Titanio,1199\n"
    )
    buf = io.BytesIO(raw)
    df = parse_uploaded_document(buf, "csv")
    assert list(df.columns) == ["descrizione", "prezzo"]
    assert "—" in df.iloc[0]["descrizione"]
    assert int(df.iloc[0]["prezzo"]) == 549


def test_csv_cp1252_italian_accents() -> None:
    """Accenti italiani in cp1252 (`è`=0xE8, `à`=0xE0): parse OK via fallback."""
    raw = (
        b"descrizione,prezzo\n"
        b"Caff\xe8 italiano,3.50\n"  # cp1252 0xE8 = 'è'
        b"Gi\xe0 pagato,5.00\n"  # cp1252 0xE0 = 'à'
    )
    buf = io.BytesIO(raw)
    df = parse_uploaded_document(buf, "csv")
    assert df.iloc[0]["descrizione"] == "Caffè italiano"
    assert df.iloc[1]["descrizione"] == "Già pagato"


def test_csv_pure_ascii_still_works() -> None:
    """Backwards-compat sentinel: ASCII puro deve continuare a funzionare."""
    raw = b"descrizione,prezzo\nGalaxy S24,549\n"
    buf = io.BytesIO(raw)
    df = parse_uploaded_document(buf, "csv")
    assert list(df.columns) == ["descrizione", "prezzo"]
    assert df.iloc[0]["descrizione"] == "Galaxy S24"


def test_xlsx_dispatch() -> None:
    from openpyxl import Workbook  # noqa: PLC0415

    wb = Workbook()
    ws = wb.active
    assert ws is not None
    ws.append(["descrizione", "prezzo"])
    ws.append(["Galaxy S24", 549])
    ws.append(["Galaxy A54", 329])
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    df = parse_uploaded_document(buf, "xlsx")
    assert list(df.columns) == ["descrizione", "prezzo"]
    assert len(df) == 2
    assert df.iloc[0]["descrizione"] == "Galaxy S24"
    assert int(df.iloc[0]["prezzo"]) == 549


def test_docx_dispatch_with_table() -> None:
    import docx  # noqa: PLC0415

    document = docx.Document()
    table = document.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "descrizione"
    table.rows[0].cells[1].text = "prezzo"
    table.rows[1].cells[0].text = "Galaxy S24"
    table.rows[1].cells[1].text = "549"
    table.rows[2].cells[0].text = "Galaxy A54"
    table.rows[2].cells[1].text = "329"
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    df = parse_uploaded_document(buf, "docx")
    assert list(df.columns) == ["descrizione", "prezzo"]
    assert len(df) == 2


def test_docx_no_matching_table_raises() -> None:
    import docx  # noqa: PLC0415

    document = docx.Document()
    document.add_paragraph("Solo testo, no tabelle.")
    buf = io.BytesIO()
    document.save(buf)
    buf.seek(0)
    with pytest.raises(ValueError, match="DOCX senza tabelle"):
        parse_uploaded_document(buf, "docx")


def test_supported_suffixes_includes_all_formats() -> None:
    assert {"csv", "xlsx", "pdf", "docx"}.issubset(SUPPORTED_SUFFIXES)
